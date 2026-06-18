from pathlib import Path

from docx import Document
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import HRFlowable, Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "outputs" / "proposals"
SOURCE = OUT / "kp_10.docx"
LOGO = ROOT / "assets" / "logo.png"
STAMP = ROOT / "assets" / "signature_stamp.png"

FONT = "Helvetica"
font_path = Path("C:/Windows/Fonts/arial.ttf")
if font_path.exists():
    pdfmetrics.registerFont(TTFont("KpFont", str(font_path)))
    FONT = "KpFont"


source_doc = Document(SOURCE)
company = source_doc.tables[0].cell(0, 1).text.strip()
hero = source_doc.tables[2].cell(0, 0).text.strip().split("\n")
number_line = hero[0] if hero else "Коммерческое предложение № 10"
date_line = hero[1] if len(hero) > 1 else "от 06.06.2026"
order_title = source_doc.tables[3].cell(0, 0).text.strip().replace("Заказ: ", "")
rows = [[cell.text.strip() for cell in row.cells] for row in source_doc.tables[4].rows[1:]]
total_lines = source_doc.tables[5].cell(0, 1).text.strip().split("\n")
total = total_lines[0].replace("Итого: ", "")
vat = total_lines[1].replace("В том числе НДС, 5%: ", "")
details = source_doc.tables[6].cell(0, 0).text.strip()
signer = source_doc.tables[7].cell(0, 0).text.strip()


def style(name, size=9, leading=None, color="#1C2730", align=TA_LEFT):
    return ParagraphStyle(
        name,
        fontName=FONT,
        fontSize=size,
        leading=leading or size * 1.25,
        textColor=colors.HexColor(color),
        alignment=align,
    )


def p(text, paragraph_style):
    return Paragraph(str(text).replace("\n", "<br/>"), paragraph_style)


def logo(width=4.2 * cm):
    if LOGO.exists():
        return Image(str(LOGO), width=width, height=width * 0.34, kind="proportional")
    return p("Mr. Puff", style("logo_fallback", 18, color="#9B49B4"))


def stamp(width=4.8 * cm):
    if STAMP.exists():
        return Image(str(STAMP), width=width, height=width * 0.78, kind="proportional")
    return ""


def build_doc(path, story, left=1.35, right=1.35, top=1.15, bottom=1.1):
    SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=left * cm,
        rightMargin=right * cm,
        topMargin=top * cm,
        bottomMargin=bottom * cm,
    ).build(story)


def items_table(header="#9B49B4", grid="#D9E0E6", zebra=True, header_white=True):
    table = Table(
        [["№", "Позиция", "Размер", "Кол-во", "НДС, 5%", "Стоимость"]] + rows,
        colWidths=[0.75 * cm, 6.3 * cm, 2.85 * cm, 1.5 * cm, 2.25 * cm, 3 * cm],
        repeatRows=1,
    )
    table_style = [
        ("FONTNAME", (0, 0), (-1, -1), FONT),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(header)),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white if header_white else colors.HexColor("#1C2730")),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor(grid)),
        ("ALIGN", (0, 0), (0, -1), "CENTER"),
        ("ALIGN", (3, 0), (-1, -1), "RIGHT"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("FONTSIZE", (0, 0), (-1, -1), 8.7),
        ("TOPPADDING", (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
    ]
    if zebra:
        table_style.append(("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F7F8FA")]))
    table.setStyle(TableStyle(table_style))
    return table


def requisites(bg="#F7F8FA", border="#E2E8D8", accent="#9B49B4"):
    box = Table([[p(details, style(f"details_{bg}", 7.5, 9.2))]], colWidths=[16.9 * cm])
    box.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(bg)),
        ("BOX", (0, 0), (-1, -1), 0.45, colors.HexColor(border)),
        ("PADDING", (0, 0), (-1, -1), 7),
    ]))
    return [p("Реквизиты", style(f"requisites_title_{bg}", 9.7, color=accent)), Spacer(1, 0.07 * cm), box]


def signature():
    table = Table([[p(signer, style("signature", 9)), stamp()]], colWidths=[8.8 * cm, 7.0 * cm])
    table.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))
    return table


def catalog():
    story = [
        Table([[logo(), p(company, style("catalog_company", 8.2, 10.5, align=TA_RIGHT))]], colWidths=[5.2 * cm, 11.7 * cm], style=[("VALIGN", (0, 0), (-1, -1), "TOP")]),
        Spacer(1, 0.15 * cm),
    ]
    strip = Table([[p("Производство мягкой мебели, подушек и текстиля на заказ", style("catalog_strip", 9, color="#6AAE4B", align=TA_CENTER))]], colWidths=[16.9 * cm])
    strip.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#EFF8EA")), ("PADDING", (0, 0), (-1, -1), 6)]))
    hero_table = Table(
        [[p(number_line + "<br/>" + date_line, style("catalog_title", 19, 24, color="#9B49B4")), p("misterpufik.ru<br/>Итого: " + total, style("catalog_total_head", 11, 18, color="#6AAE4B", align=TA_RIGHT))]],
        colWidths=[10 * cm, 6.9 * cm],
    )
    meta = Table([[p("Заказ: " + order_title, style("catalog_meta", 9))]], colWidths=[16.9 * cm])
    meta.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#EFF8EA")), ("BOX", (0, 0), (-1, -1), 0.45, colors.HexColor("#DDEFD6")), ("PADDING", (0, 0), (-1, -1), 8)]))
    total_box = Table([["", p("Итого: " + total + "<br/>В том числе НДС, 5%: " + vat, style("catalog_total", 12, 15, color="#9B49B4", align=TA_RIGHT))]], colWidths=[8.2 * cm, 8.7 * cm])
    total_box.setStyle(TableStyle([("BACKGROUND", (1, 0), (1, 0), colors.HexColor("#F8F2FC")), ("BOX", (1, 0), (1, 0), 0.45, colors.HexColor("#E7D8EF")), ("PADDING", (1, 0), (1, 0), 10)]))
    story += [strip, Spacer(1, 0.35 * cm), hero_table, Spacer(1, 0.3 * cm), meta, Spacer(1, 0.3 * cm), items_table(), Spacer(1, 0.25 * cm), total_box, Spacer(1, 0.2 * cm)]
    story += requisites() + [Spacer(1, 0.2 * cm), signature()]
    build_doc(OUT / "kp_10_variant_1_firmenny_catalog.pdf", story)


def premium():
    story = [
        Table([[logo(3.8 * cm), p(company, style("premium_company", 7.8, 9.8, "#3A3A3A", TA_RIGHT))]], colWidths=[4.7 * cm, 11.6 * cm], style=[("VALIGN", (0, 0), (-1, -1), "TOP")]),
        Spacer(1, 0.25 * cm),
        HRFlowable(width="100%", thickness=1.2, color=colors.HexColor("#9B49B4")),
        Spacer(1, 0.45 * cm),
        p(number_line, style("premium_title", 21, 25)),
        p(date_line, style("premium_date", 9, color="#6B7680")),
        Spacer(1, 0.25 * cm),
    ]
    chips = Table([[p("Индивидуальное производство", style("premium_chip_1", 8, color="#6AAE4B", align=TA_CENTER)), p("Подушки и текстиль", style("premium_chip_2", 8, color="#9B49B4", align=TA_CENTER)), p(total, style("premium_chip_3", 12, align=TA_RIGHT))]], colWidths=[5.2 * cm, 5.2 * cm, 6.1 * cm])
    chips.setStyle(TableStyle([("LINEBELOW", (0, 0), (-1, -1), 0.3, colors.HexColor("#E6E9EC")), ("TOPPADDING", (0, 0), (-1, -1), 6), ("BOTTOMPADDING", (0, 0), (-1, -1), 7)]))
    table = items_table("#FFFFFF", "#D8DEE6", False, False)
    table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), FONT),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1C2730")),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D8DEE6")),
        ("LINEBELOW", (0, 0), (-1, 0), 1.0, colors.HexColor("#9B49B4")),
        ("ALIGN", (0, 0), (0, -1), "CENTER"),
        ("ALIGN", (3, 0), (-1, -1), "RIGHT"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("FONTSIZE", (0, 0), (-1, -1), 8.7),
        ("PADDING", (0, 0), (-1, -1), 7),
    ]))
    total_row = Table([["", p("Итого: " + total + "<br/>НДС, 5%: " + vat, style("premium_total", 14, 18, align=TA_RIGHT))]], colWidths=[9.6 * cm, 6.7 * cm])
    total_row.setStyle(TableStyle([("LINEABOVE", (1, 0), (1, 0), 1.2, colors.HexColor("#6AAE4B")), ("TOPPADDING", (1, 0), (1, 0), 8)]))
    story += [chips, Spacer(1, 0.35 * cm), p("Заказ: " + order_title, style("premium_meta", 9.2)), Spacer(1, 0.18 * cm), table, Spacer(1, 0.25 * cm), total_row, Spacer(1, 0.25 * cm)]
    story += requisites("#FFFFFF", "#D8DEE6", "#1C2730") + [Spacer(1, 0.15 * cm), signature()]
    build_doc(OUT / "kp_10_variant_2_premium_minimal.pdf", story, left=1.55, right=1.55, top=1.3)


def production():
    band = Table([[logo(4.1 * cm), p(number_line + "<br/>" + date_line, style("production_header", 17, 21, "#FFFFFF")), p(total + "<br/>итог по КП", style("production_head_total", 16, 19, "#FFFFFF", TA_RIGHT))]], colWidths=[4.6 * cm, 8.1 * cm, 4.2 * cm])
    band.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#6AAE4B")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("PADDING", (0, 0), (-1, -1), 11)]))
    meta = Table([[p("Объект расчёта<br/>" + order_title, style("production_meta_1", 9.3)), p("Формат<br/>производство под заказ", style("production_meta_2", 9.3)), p("НДС<br/>" + vat, style("production_meta_3", 9.3, align=TA_RIGHT))]], colWidths=[7.2 * cm, 5.2 * cm, 3.7 * cm])
    meta.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F7F8FA")), ("BOX", (0, 0), (-1, -1), 0.45, colors.HexColor("#D9E0E6")), ("INNERGRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#D9E0E6")), ("PADDING", (0, 0), (-1, -1), 8)]))
    total_grid = Table([[p("Итого к оплате", style("production_total_label", 9, color="#6B7680")), p(total, style("production_total_value", 16, color="#6AAE4B", align=TA_RIGHT))], [p("В том числе НДС, 5%", style("production_vat_label", 8, color="#6B7680")), p(vat, style("production_vat", 9, align=TA_RIGHT))]], colWidths=[8 * cm, 8.1 * cm])
    total_grid.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#EFF8EA")), ("BOX", (0, 0), (-1, -1), 0.45, colors.HexColor("#DDEFD6")), ("PADDING", (0, 0), (-1, -1), 8)]))
    story = [band, Spacer(1, 0.22 * cm), p(company, style("production_company", 7.6, 9.4, "#4A4F55", TA_RIGHT)), Spacer(1, 0.28 * cm), meta, Spacer(1, 0.3 * cm), items_table("#6AAE4B"), Spacer(1, 0.22 * cm), total_grid, Spacer(1, 0.22 * cm)]
    story += requisites("#F7F8FA", "#D9E0E6", "#6AAE4B") + [Spacer(1, 0.16 * cm), signature()]
    build_doc(OUT / "kp_10_variant_3_proizvodstvenny.pdf", story, left=1.3, right=1.3, top=1.05)


def bright():
    hero_table = Table([[logo(4.5 * cm), p(number_line + "<br/>" + date_line + "<br/><br/>" + total, style("bright_header", 18, 22, "#FFFFFF", TA_RIGHT))]], colWidths=[5.3 * cm, 11.6 * cm])
    hero_table.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#9B49B4")), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("PADDING", (0, 0), (-1, -1), 14)]))
    chips = Table([[p("Гарантия<br/>12 месяцев", style("bright_chip_1", 8.5, 10.5, "#6AAE4B", TA_CENTER)), p("Доставка<br/>по России", style("bright_chip_2", 8.5, 10.5, "#6AAE4B", TA_CENTER)), p("Производство<br/>СПб", style("bright_chip_3", 8.5, 10.5, "#6AAE4B", TA_CENTER))]], colWidths=[5.63 * cm] * 3)
    chips.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#EFF8EA")), ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.white), ("PADDING", (0, 0), (-1, -1), 8)]))
    meta = Table([[p("Заказ: " + order_title, style("bright_meta", 9.5, color="#9B49B4"))]], colWidths=[16.9 * cm])
    meta.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F8F2FC")), ("BOX", (0, 0), (-1, -1), 0.45, colors.HexColor("#E7D8EF")), ("PADDING", (0, 0), (-1, -1), 9)]))
    total_banner = Table([[p("Итого: " + total, style("bright_total", 17, color="#FFFFFF")), p("НДС, 5%: " + vat, style("bright_vat", 9, color="#FFFFFF", align=TA_RIGHT))]], colWidths=[9.4 * cm, 7.5 * cm])
    total_banner.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#6AAE4B")), ("PADDING", (0, 0), (-1, -1), 10), ("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))
    story = [hero_table, chips, Spacer(1, 0.26 * cm), p(company, style("bright_company", 7.5, 9.2, "#4A4F55")), Spacer(1, 0.2 * cm), meta, Spacer(1, 0.28 * cm), items_table("#9B49B4", "#E3DAEA", True), Spacer(1, 0.25 * cm), total_banner, Spacer(1, 0.23 * cm)]
    story += requisites("#FFFFFF", "#E3DAEA", "#9B49B4") + [Spacer(1, 0.13 * cm), signature()]
    build_doc(OUT / "kp_10_variant_4_yarkiy_klientskiy.pdf", story, left=1.25, right=1.25, top=1.0)


def main():
    catalog()
    premium()
    production()
    bright()
    for path in sorted(OUT.glob("kp_10_variant_*.pdf")):
        print(path.resolve())


if __name__ == "__main__":
    main()
