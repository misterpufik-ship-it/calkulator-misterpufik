from __future__ import annotations

import base64
import html
import json
import mimetypes
import os
import re
import tempfile
import urllib.error
import urllib.request
from datetime import datetime
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.parse import parse_qs, unquote, urlsplit

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from avito_telegram_bridge import (
    BridgeError,
    bridge_status,
    check_yandex_mail,
    handle_avito_webhook,
    handle_telegram_webhook,
    handle_whatsapp_webhook,
    verify_whatsapp_subscription,
)

ROOT = Path(__file__).resolve().parent
OUTPUTS = ROOT / "outputs" / "proposals"
OUTPUTS.mkdir(parents=True, exist_ok=True)
LOGO_PATH = ROOT / "assets" / "logo.png"
SIGNATURE_STAMP_PATH = ROOT / "assets" / "signature_stamp.png"
BLACK = RGBColor(20, 20, 20)
PURPLE = RGBColor(155, 73, 180)
GREEN = RGBColor(106, 174, 75)
SOFT_PURPLE = RGBColor(248, 242, 252)
SOFT_GREEN = RGBColor(239, 248, 234)
SOFT_GRAY = RGBColor(247, 248, 250)
BRAND_PURPLE_HEX = "9B49B4"
BRAND_GREEN_HEX = "6AAE4B"
SOFT_PURPLE_HEX = "F8F2FC"
SOFT_GREEN_HEX = "EFF8EA"
SOFT_GRAY_HEX = "F7F8FA"


def safe_name(value: str) -> str:
    value = re.sub(r"[^0-9A-Za-zА-Яа-я_-]+", "_", value.strip())
    return value.strip("_") or "proposal"


def money(value: float) -> str:
    return f"{round(float(value)):,.0f}".replace(",", " ") + " ₽"


def production_term_text(value: str) -> str:
    value = str(value or "").strip()
    if not value:
        return ""
    suffix = "с момента поступления оплаты"
    lower = value.lower()
    if suffix in lower:
        return value
    if "рабоч" in lower:
        return f"{value} {suffix}"
    return f"{value} рабочих дней {suffix}"


def decode_data_url(data_url: str) -> Path | None:
    if not data_url or "," not in data_url:
        return None
    meta, raw = data_url.split(",", 1)
    suffix = ".png"
    if "jpeg" in meta or "jpg" in meta:
        suffix = ".jpg"
    if "webp" in meta:
        suffix = ".webp"
    path = Path(tempfile.NamedTemporaryFile(delete=False, suffix=suffix).name)
    path.write_bytes(base64.b64decode(raw))
    return path


def image_or_fallback(data_url: str, fallback: Path | None = None) -> Path | None:
    decoded = decode_data_url(data_url)
    if decoded:
        return decoded
    if fallback and fallback.exists():
        return fallback
    return None


def set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_table_borders(table, color: str = "DADDE3"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 120, start: int = 140, bottom: int = 120, end: int = 140):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def style_doc(doc: Document):
    section = doc.sections[0]
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)


def add_doc_paragraph(doc: Document, text: str, size: int = 10, bold: bool = False, align=None, color=None):
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return paragraph


def build_docx(payload: dict, output: Path):
    doc = Document()
    style_doc(doc)
    company = payload.get("company", {})

    header = doc.add_table(rows=1, cols=2)
    header.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(header, "FFFFFF")
    header.columns[0].width = Cm(5.6)
    header.columns[1].width = Cm(11.4)
    for cell in header.rows[0].cells:
        set_cell_margins(cell, 80, 80, 120, 80)
    header.cell(0, 0).text = ""
    if LOGO_PATH.exists():
        header.cell(0, 0).paragraphs[0].add_run().add_picture(str(LOGO_PATH), width=Cm(4.6))
    else:
        set_cell_text(header.cell(0, 0), "Mr. Puff", True)
    set_cell_text(header.cell(0, 1), company.get("companyName") or "Компания", True)
    header.cell(0, 1).paragraphs[0].runs[0].font.size = Pt(11)
    header.cell(0, 1).paragraphs[0].runs[0].font.color.rgb = PURPLE
    contacts = "\n".join(filter(None, [
        company.get("companyPhone", ""),
        company.get("companyEmail", ""),
        company.get("companySite", ""),
        company.get("companyAddress", ""),
        "Юридический адрес: 197349, Санкт-Петербург, проспект Сизова, д. 34/18, кв. 452",
        "ИНН: 781432056933",
        "ОГРНИП: 312784734000451",
    ]))
    if contacts:
        p = header.cell(0, 1).add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(contacts)
        r.font.name = "Arial"
        r.font.size = Pt(8.5)
    header.cell(0, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    brand_strip = doc.add_table(rows=1, cols=1)
    brand_strip.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(brand_strip, "FFFFFF")
    shade_cell(brand_strip.cell(0, 0), SOFT_GREEN_HEX)
    set_cell_margins(brand_strip.cell(0, 0), 70, 180, 70, 180)
    p = brand_strip.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Производство мягкой мебели, подушек и текстиля на заказ")
    r.font.name = "Arial"
    r.font.size = Pt(9)
    r.bold = True
    r.font.color.rgb = GREEN

    hero = doc.add_table(rows=1, cols=2)
    hero.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(hero, "FFFFFF")
    hero.columns[0].width = Cm(11)
    hero.columns[1].width = Cm(6)
    left = hero.cell(0, 0)
    right_cell = hero.cell(0, 1)
    for cell in hero.rows[0].cells:
        set_cell_margins(cell, 220, 80, 180, 80)
    left.text = ""
    p = left.paragraphs[0]
    r = p.add_run(f"Коммерческое предложение № {payload['number']}")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(20)
    r.font.color.rgb = PURPLE
    p2 = left.add_paragraph()
    r2 = p2.add_run(f"от {payload['date']}")
    r2.font.name = "Arial"
    r2.font.size = Pt(9)
    r2.font.color.rgb = BLACK
    right_cell.text = ""
    p = right_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(company.get("companySite") or "misterpufik.ru")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(11)
    r.font.color.rgb = GREEN
    p_total = right_cell.add_paragraph()
    p_total.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    total_label = p_total.add_run(f"Итого: {money(payload['total'])}")
    total_label.bold = True
    total_label.font.name = "Arial"
    total_label.font.size = Pt(13)
    total_label.font.color.rgb = PURPLE

    meta = doc.add_table(rows=1, cols=1)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(meta, "DDEFD6")
    shade_cell(meta.cell(0, 0), SOFT_GREEN_HEX)
    set_cell_margins(meta.cell(0, 0), 150, 220, 150, 220)
    set_cell_text(meta.cell(0, 0), f"Заказ: {payload.get('title') or 'Заказ подушек'}", True)

    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table, "D9DEE8")
    widths = [Cm(1), Cm(6.4), Cm(3), Cm(1.7), Cm(2.4), Cm(3)]
    headers = ["№", "Позиция", "Размер", "Кол-во", "НДС, 5%", "Стоимость"]
    for i, text in enumerate(headers):
        set_cell_text(table.cell(0, i), text, True)
        table.cell(0, i).width = widths[i]
        shade_cell(table.cell(0, i), BRAND_PURPLE_HEX)
        set_cell_margins(table.cell(0, i), 120, 120, 120, 120)
        table.cell(0, i).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    for row_index, line in enumerate(payload.get("lines", []), 1):
        row = table.add_row()
        values = [line["index"], line["type"], line["size"], line["quantity"], money(line.get("vatAmount", 0)), money(line["totalPrice"])]
        for i, value in enumerate(values):
            set_cell_text(row.cells[i], value)
            row.cells[i].width = widths[i]
            set_cell_margins(row.cells[i], 120, 120, 120, 120)
            if row_index % 2 == 0:
                shade_cell(row.cells[i], SOFT_GRAY_HEX)
    if float(payload.get("deliveryGrossAmount") or 0) > 0:
        row = table.add_row()
        values = [len(payload.get("lines", [])) + 1, "Доставка", "", 1, money(payload.get("deliveryVat", 0)), money(payload["deliveryGrossAmount"])]
        for i, value in enumerate(values):
            set_cell_text(row.cells[i], value, i == 1)
            row.cells[i].width = widths[i]
            set_cell_margins(row.cells[i], 120, 120, 120, 120)
            shade_cell(row.cells[i], SOFT_GREEN_HEX)
            if i == 1:
                row.cells[i].paragraphs[0].runs[0].font.color.rgb = GREEN

    total_box = doc.add_table(rows=1, cols=2)
    total_box.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(total_box, "FFFFFF")
    total_box.columns[0].width = Cm(9)
    total_box.columns[1].width = Cm(8)
    total_box.cell(0, 0).text = ""
    total_box.cell(0, 1).text = ""
    shade_cell(total_box.cell(0, 1), SOFT_PURPLE_HEX)
    set_cell_margins(total_box.cell(0, 1), 180, 220, 180, 220)
    p = total_box.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(f"Итого: {money(payload['total'])}")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(15)
    r.font.color.rgb = PURPLE
    if float(payload.get("totalVat") or 0) > 0:
        vat_p = total_box.cell(0, 1).add_paragraph()
        vat_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        vat_p.paragraph_format.space_before = Pt(5)
        vat_r = vat_p.add_run(f"В том числе НДС, 5%: {money(payload['totalVat'])}")
        vat_r.font.name = "Arial"
        vat_r.font.size = Pt(8)
        vat_r.font.color.rgb = BLACK
    if payload.get("productionTerm"):
        term = add_doc_paragraph(doc, f"Срок изготовления: {production_term_text(payload['productionTerm'])}", 9, True)
        term.paragraph_format.space_before = Pt(8)

    details = str(company.get("companyDetails") or "").strip()
    if details:
        details_title = add_doc_paragraph(doc, "Реквизиты", 10, True, color=(155, 73, 180))
        details_title.paragraph_format.space_before = Pt(8)
        details_box = doc.add_table(rows=1, cols=1)
        set_table_borders(details_box, "E2E8D8")
        shade_cell(details_box.cell(0, 0), SOFT_GRAY_HEX)
        set_cell_margins(details_box.cell(0, 0), 130, 170, 130, 170)
        details_box.cell(0, 0).text = ""
        p_details = details_box.cell(0, 0).paragraphs[0]
        r_details = p_details.add_run(details)
        r_details.font.name = "Arial"
        r_details.font.size = Pt(8)
        r_details.font.color.rgb = BLACK

    sign_table = doc.add_table(rows=1, cols=2)
    sign_table.columns[0].width = Cm(9)
    sign_table.columns[1].width = Cm(8)
    set_table_borders(sign_table, "FFFFFF")
    set_cell_text(sign_table.cell(0, 0), company.get("signerTitle") or "ИП")
    sign_table.cell(0, 0).paragraphs[0].add_run("   " + (company.get("signerName") or ""))
    sign_table.cell(0, 1).text = ""
    stamp = image_or_fallback(company.get("stampDataUrl", ""), SIGNATURE_STAMP_PATH)
    if stamp:
        sign_table.cell(0, 1).paragraphs[0].add_run().add_picture(str(stamp), width=Cm(5.2))

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = f"{company.get('companyName') or 'Компания'} | Коммерческое предложение № {payload['number']}"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.name = "Arial"
    doc.save(output)


def register_pdf_font() -> str:
    candidates = [
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("C:/Windows/Fonts/calibri.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            pdfmetrics.registerFont(TTFont("ProposalFont", str(candidate)))
            return "ProposalFont"
    return "Helvetica"


def build_pdf(payload: dict, output: Path):
    font = register_pdf_font()
    styles = getSampleStyleSheet()
    normal = ParagraphStyle("normal_ru", parent=styles["Normal"], fontName=font, fontSize=9, leading=12, textColor=colors.HexColor("#1c2730"))
    right = ParagraphStyle("right_ru", parent=normal, alignment=TA_RIGHT, fontSize=15, leading=18, textColor=colors.HexColor("#9B49B4"))
    company = payload.get("company", {})
    doc = SimpleDocTemplate(str(output), pagesize=A4, rightMargin=1.45 * cm, leftMargin=1.45 * cm, topMargin=1.2 * cm, bottomMargin=1.2 * cm)
    story = []

    contacts = "<br/>".join(filter(None, [
        company.get("companyPhone", ""),
        company.get("companyEmail", ""),
        company.get("companySite", ""),
        company.get("companyAddress", ""),
        "Юридический адрес: 197349, Санкт-Петербург, проспект Сизова, д. 34/18, кв. 452",
        "ИНН: 781432056933",
        "ОГРНИП: 312784734000451",
    ]))
    logo_cell = Image(str(LOGO_PATH), width=4.6 * cm, height=1.48 * cm, kind="proportional") if LOGO_PATH.exists() else Paragraph("<b>Mr. Puff</b>", normal)
    brand_style = ParagraphStyle("brand", parent=normal, fontName=font, fontSize=8.5, leading=11, textColor=colors.HexColor("#141414"), alignment=TA_RIGHT)
    header_table = Table([[logo_cell, Paragraph(f"<b>{company.get('companyName') or 'Компания'}</b><br/>{contacts}", brand_style)]], colWidths=[5.8 * cm, 11.1 * cm])
    header_table.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP")]))
    story.append(header_table)
    story.append(Spacer(1, 0.18 * cm))

    strip = Table([[Paragraph("<b>Производство мягкой мебели, подушек и текстиля на заказ</b>", ParagraphStyle("strip", parent=normal, alignment=TA_CENTER, textColor=colors.HexColor("#6AAE4B"), fontName=font, fontSize=9, leading=11))]], colWidths=[16.9 * cm])
    strip.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#EFF8EA")),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(strip)
    story.append(Spacer(1, 0.35 * cm))

    hero_left = Paragraph(
        f"<font color='#9B49B4'><b>Коммерческое предложение № {payload['number']}</b></font><br/><font size='9'>от {payload['date']}</font>",
        ParagraphStyle("hero_left", parent=normal, fontName=font, fontSize=18, leading=22),
    )
    hero_right = Paragraph(
        f"<font color='#6AAE4B'><b>{company.get('companySite') or 'misterpufik.ru'}</b></font><br/><font color='#9B49B4' size='13'><b>Итого: {money(payload['total'])}</b></font>",
        ParagraphStyle("hero_right", parent=normal, fontName=font, fontSize=11, leading=17, alignment=TA_RIGHT),
    )
    hero_table = Table([[hero_left, hero_right]], colWidths=[8.5 * cm, 8.4 * cm])
    hero_table.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))
    story.append(hero_table)
    story.append(Spacer(1, 0.35 * cm))

    meta_data = [[Paragraph(f"<b>Заказ:</b> {payload.get('title') or 'Заказ подушек'}", normal)]]
    meta_table = Table(meta_data, colWidths=[16.9 * cm])
    meta_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#EFF8EA")),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#DDEFD6")),
        ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#DDEFD6")),
        ("LEFTPADDING", (0, 0), (-1, -1), 10),
        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 0.35 * cm))

    data = [["№", "Позиция", "Размер", "Кол-во", "НДС, 5%", "Стоимость"]]
    for line in payload.get("lines", []):
        data.append([line["index"], line["type"], line["size"], line["quantity"], money(line.get("vatAmount", 0)), money(line["totalPrice"])])
    if float(payload.get("deliveryGrossAmount") or 0) > 0:
        data.append([len(payload.get("lines", [])) + 1, "Доставка", "", 1, money(payload.get("deliveryVat", 0)), money(payload["deliveryGrossAmount"])])
    table = Table(data, colWidths=[0.85 * cm, 6.2 * cm, 2.85 * cm, 1.55 * cm, 2.25 * cm, 3.05 * cm])
    table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), font),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#9B49B4")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#d9e0e6")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F7F8FA")]),
        ("ALIGN", (0, 0), (0, -1), "CENTER"),
        ("ALIGN", (3, 0), (-1, -1), "RIGHT"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("FONTNAME", (0, 0), (-1, 0), font),
        ("FONTSIZE", (0, 0), (-1, 0), 9),
        ("FONTSIZE", (0, 1), (-1, -1), 9),
    ]))
    story.append(table)
    story.append(Spacer(1, 0.32 * cm))
    total_rows = [[Paragraph(f"<b>Итого: {money(payload['total'])}</b>", right)]]
    if float(payload.get("totalVat") or 0) > 0:
        vat_style = ParagraphStyle("vat_right", parent=normal, alignment=TA_RIGHT, fontSize=8, leading=10, fontName=font)
        total_rows.append([Paragraph(f"В том числе НДС, 5%: {money(payload['totalVat'])}", vat_style)])
    total_table = Table([["", Table(total_rows, colWidths=[7.1 * cm])]], colWidths=[9.8 * cm, 7.1 * cm])
    total_table.setStyle(TableStyle([
        ("BACKGROUND", (1, 0), (1, 0), colors.HexColor("#F8F2FC")),
        ("BOX", (1, 0), (1, 0), 0.5, colors.HexColor("#E7D8EF")),
        ("LEFTPADDING", (1, 0), (1, 0), 12),
        ("RIGHTPADDING", (1, 0), (1, 0), 12),
        ("TOPPADDING", (1, 0), (1, 0), 8),
        ("BOTTOMPADDING", (1, 0), (1, 0), 8),
    ]))
    story.append(total_table)
    if payload.get("productionTerm"):
        story.append(Spacer(1, 0.22 * cm))
        story.append(Paragraph(f"<b>Срок изготовления:</b> {production_term_text(payload['productionTerm'])}", normal))
    details = str(company.get("companyDetails") or "").strip()
    if details:
        story.append(Spacer(1, 0.22 * cm))
        story.append(Paragraph("<font color='#9B49B4'><b>Реквизиты</b></font>", normal))
        details_table = Table([[Paragraph(details.replace("\n", "<br/>"), ParagraphStyle("details", parent=normal, fontSize=8, leading=10, fontName=font))]], colWidths=[16.9 * cm])
        details_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F7F8FA")),
            ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#E2E8D8")),
            ("LEFTPADDING", (0, 0), (-1, -1), 9),
            ("RIGHTPADDING", (0, 0), (-1, -1), 9),
            ("TOPPADDING", (0, 0), (-1, -1), 7),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ]))
        story.append(details_table)
    stamp = image_or_fallback(company.get("stampDataUrl", ""), SIGNATURE_STAMP_PATH)
    sign_cells = [
        Paragraph(f"<b>{company.get('signerTitle') or 'ИП'}</b>&nbsp;&nbsp;&nbsp;{company.get('signerName') or ''}", normal),
        "",
    ]
    if stamp:
        sign_cells[1] = Image(str(stamp), width=5.2 * cm, height=4.2 * cm, kind="proportional")
    story.append(Spacer(1, 0.35 * cm))
    sign_table_pdf = Table([sign_cells], colWidths=[8.7 * cm, 7.2 * cm])
    sign_table_pdf.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "MIDDLE")]))
    story.append(sign_table_pdf)
    doc.build(story)


def build_company_card_pdf(payload: dict, output: Path):
    font = register_pdf_font()
    styles = getSampleStyleSheet()
    company = payload.get("company", {})
    normal = ParagraphStyle("card_normal", parent=styles["Normal"], fontName=font, fontSize=10, leading=14, textColor=colors.HexColor("#243C32"))
    label_style = ParagraphStyle("card_label", parent=normal, fontSize=8, leading=11, textColor=colors.HexColor("#6F7B73"))
    value_style = ParagraphStyle("card_value", parent=normal, fontSize=10, leading=14, textColor=colors.HexColor("#243C32"))
    title_style = ParagraphStyle("card_title", parent=normal, alignment=TA_CENTER, fontSize=30, leading=34, textColor=colors.white)
    subtitle_style = ParagraphStyle("card_subtitle", parent=normal, alignment=TA_CENTER, fontSize=13, leading=16, textColor=colors.white)

    doc = SimpleDocTemplate(str(output), pagesize=A4, rightMargin=1.45 * cm, leftMargin=1.45 * cm, topMargin=1.25 * cm, bottomMargin=1.25 * cm)
    story = []

    logo = Paragraph("M I S T E R<br/>P U F I K", title_style)
    b2b = Paragraph("&middot; B 2 B &middot;", subtitle_style)
    logo_box = Table([[logo], [b2b]], colWidths=[16.9 * cm])
    logo_box.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#395648")),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("TOPPADDING", (0, 0), (-1, 0), 36),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 4),
        ("TOPPADDING", (0, 1), (-1, 1), 0),
        ("BOTTOMPADDING", (0, 1), (-1, 1), 34),
    ]))
    story.append(logo_box)
    story.append(Spacer(1, 0.5 * cm))

    heading = Paragraph(f"<b>{html.escape(company.get('companyName') or 'Компания')}</b>", ParagraphStyle("card_heading", parent=normal, fontName=font, fontSize=17, leading=22, textColor=colors.HexColor("#243C32")))
    story.append(heading)
    story.append(Spacer(1, 0.28 * cm))

    contact_rows = [
        ("Телефон", company.get("companyPhone", "")),
        ("Email", company.get("companyEmail", "")),
        ("Сайт", company.get("companySite", "")),
        ("Адрес", company.get("companyAddress", "")),
    ]
    detail_text = str(company.get("companyDetails") or "").strip()
    for line in detail_text.splitlines():
        if ":" in line:
            label, value = line.split(":", 1)
            contact_rows.append((label.strip(), value.strip()))
        elif line.strip():
            contact_rows.append(("Реквизиты", line.strip()))

    data = [[Paragraph(html.escape(label), label_style), Paragraph(html.escape(str(value or "-")), value_style)] for label, value in contact_rows]
    table = Table(data, colWidths=[4.2 * cm, 12.7 * cm])
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#D9DED8")),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F2F6F1")),
        ("ROWBACKGROUNDS", (1, 0), (1, -1), [colors.white, colors.HexColor("#F7FAF6")]),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 9),
        ("RIGHTPADDING", (0, 0), (-1, -1), 9),
        ("TOPPADDING", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(table)
    story.append(Spacer(1, 0.35 * cm))
    story.append(Paragraph("Карточка компании сформирована без печати и подписи.", ParagraphStyle("card_note", parent=normal, alignment=TA_CENTER, fontSize=8, leading=10, textColor=colors.HexColor("#6F7B73"))))
    doc.build(story)


def oksana_rows(payload: dict) -> list[list[object]]:
    rows = []
    for line in payload.get("lines", []):
        quantity = int(float(line.get("quantity") or 0))
        unit = float(line.get("sewingPrice") or 0)
        total = float(line.get("sewingTotal") or (unit * quantity))
        assembly = float(line.get("assemblyTotal") or 0)
        rows.append([
            int(line.get("index") or len(rows) + 1),
            line.get("type", ""),
            line.get("size", ""),
            quantity,
            unit,
            assembly,
            total,
        ])
    return rows


def build_oksana_docx(payload: dict, output: Path):
    doc = Document()
    style_doc(doc)
    add_doc_paragraph(doc, f"Расчет Оксана № {payload['number']}", 16, True, WD_ALIGN_PARAGRAPH.CENTER, (20, 20, 20))
    add_doc_paragraph(doc, f"от {payload['date']}", 9, False, WD_ALIGN_PARAGRAPH.CENTER)
    add_doc_paragraph(doc, f"Заказ: {payload.get('title') or 'Заказ подушек'}", 10, True)

    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table, "D9DEE8")
    widths = [Cm(1), Cm(5.6), Cm(3), Cm(1.7), Cm(2.5), Cm(2.4), Cm(2.6)]
    headers = ["№", "Позиция", "Размер", "Кол-во", "Пошив за 1 шт", "Сборка", "Сумма пошива"]
    for i, header in enumerate(headers):
        set_cell_text(table.cell(0, i), header, True)
        table.cell(0, i).width = widths[i]
        shade_cell(table.cell(0, i), "F1ECFF")

    total_sum = 0
    for row_data in oksana_rows(payload):
        row = table.add_row()
        total_sum += float(row_data[6])
        values = [row_data[0], row_data[1], row_data[2], row_data[3], money(row_data[4]), money(row_data[5]), money(row_data[6])]
        for i, value in enumerate(values):
            set_cell_text(row.cells[i], value, i == 1)
            row.cells[i].width = widths[i]

    total_p = add_doc_paragraph(doc, f"Итого по позициям: {money(total_sum)}", 12, True, WD_ALIGN_PARAGRAPH.RIGHT, (138, 43, 226))
    total_p.paragraph_format.space_before = Pt(10)
    doc.save(output)


def build_oksana_xlsx(payload: dict, output: Path):
    wb = Workbook()
    ws = wb.active
    ws.title = "Расчет Оксана"

    purple = "8A2BE2"
    light = "F1ECFF"
    border = Border(
        left=Side(style="thin", color="D9DEE8"),
        right=Side(style="thin", color="D9DEE8"),
        top=Side(style="thin", color="D9DEE8"),
        bottom=Side(style="thin", color="D9DEE8"),
    )

    ws["A1"] = f"Расчет Оксана № {payload['number']}"
    ws["A1"].font = Font(name="Arial", size=16, bold=True, color=purple)
    ws.merge_cells("A1:G1")
    ws["A2"] = f"от {payload['date']}"
    ws.merge_cells("A2:G2")
    ws["A3"] = f"Заказ: {payload.get('title') or 'Заказ подушек'}"
    ws["A3"].font = Font(name="Arial", size=11, bold=True)
    ws.merge_cells("A3:G3")

    headers = ["№", "Позиция", "Размер", "Кол-во", "Пошив за 1 шт", "Сборка", "Сумма пошива"]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=5, column=col, value=header)
        cell.font = Font(name="Arial", size=10, bold=True)
        cell.fill = PatternFill("solid", fgColor=light)
        cell.border = border
        cell.alignment = Alignment(horizontal="center", vertical="center")

    total_sum = 0
    for row_index, row_data in enumerate(oksana_rows(payload), 6):
        total_sum += float(row_data[6])
        for col, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_index, column=col, value=value)
            cell.font = Font(name="Arial", size=10, bold=(col == 2))
            cell.border = border
            cell.alignment = Alignment(vertical="center", horizontal="right" if col in (4, 5, 6, 7) else "left")
            if col in (5, 6, 7):
                cell.number_format = '#,##0" ₽"'

    total_row = 6 + len(oksana_rows(payload))
    ws.cell(row=total_row, column=6, value="Итого").font = Font(name="Arial", size=11, bold=True, color=purple)
    total_cell = ws.cell(row=total_row, column=7, value=total_sum)
    total_cell.font = Font(name="Arial", size=11, bold=True, color=purple)
    total_cell.number_format = '#,##0" ₽"'

    widths = [8, 32, 18, 12, 16, 14, 16]
    for index, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(index)].width = width
    for row in range(1, total_row + 1):
        ws.row_dimensions[row].height = 24
    ws.freeze_panes = "A6"
    wb.save(output)


class Handler(SimpleHTTPRequestHandler):
    def send_json(self, payload: dict, status: int = 200):
        body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def send_text(self, text: str, status: int = 200):
        body = text.encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "text/plain; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def read_body(self) -> bytes:
        length = int(self.headers.get("Content-Length", "0"))
        if length <= 0:
            return b""
        return self.rfile.read(length)

    def read_json_payload(self) -> dict:
        raw_body = self.read_body()
        if not raw_body:
            return {}
        return json.loads(raw_body.decode("utf-8"))

    def do_GET(self):
        url = urlsplit(self.path)
        path = url.path
        if path == "/api/avito/status":
            self.send_json(bridge_status())
            return
        if path == "/api/whatsapp/webhook":
            try:
                challenge = verify_whatsapp_subscription(parse_qs(url.query))
                self.send_text(challenge)
            except BridgeError as error:
                self.send_json({"ok": False, "error": str(error)}, 403)
            return
        super().do_GET()

    def do_POST(self):
        path = self.path.split("?", 1)[0]

        if path == "/api/avito/webhook":
            try:
                result = handle_avito_webhook(dict(self.headers), self.read_json_payload())
                self.send_json(result)
            except BridgeError as error:
                self.send_json({"ok": False, "error": str(error)}, 502)
            except Exception as error:
                self.send_json({"ok": False, "error": str(error)}, 500)
            return

        if path == "/api/telegram/webhook":
            try:
                result = handle_telegram_webhook(dict(self.headers), self.read_json_payload())
                self.send_json(result)
            except BridgeError as error:
                self.send_json({"ok": False, "error": str(error)}, 400)
            except Exception as error:
                self.send_json({"ok": False, "error": str(error)}, 500)
            return

        if path == "/api/whatsapp/webhook":
            try:
                raw_body = self.read_body()
                payload = json.loads(raw_body.decode("utf-8")) if raw_body else {}
                result = handle_whatsapp_webhook(dict(self.headers), payload, raw_body)
                self.send_json(result)
            except BridgeError as error:
                self.send_json({"ok": False, "error": str(error)}, 400)
            except Exception as error:
                self.send_json({"ok": False, "error": str(error)}, 500)
            return

        if path == "/api/yandex/check":
            try:
                payload = self.read_json_payload()
                limit = int(payload.get("limit") or 10)
                self.send_json(check_yandex_mail(limit))
            except BridgeError as error:
                self.send_json({"ok": False, "error": str(error)}, 400)
            except Exception as error:
                self.send_json({"ok": False, "error": str(error)}, 500)
            return

        if path == "/api/rental":
            payload = self.read_json_payload()
            webhook_url = str(payload.get("webhookUrl") or "").strip()
            if not webhook_url.startswith("https://script.google.com/"):
                self.send_json({"ok": False, "error": "Укажите корректный URL Google Apps Script webhook."}, 400)
                return

            row_payload = {
                "date": payload.get("date", ""),
                "client": payload.get("client", ""),
                "phone": payload.get("phone", ""),
                "driver": payload.get("driver", ""),
                "item": payload.get("item", ""),
                "quantity": payload.get("quantity", ""),
                "amount": payload.get("amount", ""),
                "startDate": payload.get("startDate", ""),
                "startTime": payload.get("startTime", ""),
                "endDate": payload.get("endDate", ""),
                "endTime": payload.get("endTime", ""),
                "status": payload.get("status", ""),
                "colorSummary": payload.get("colorSummary", ""),
                "colors": payload.get("colors", []),
                "lines": payload.get("lines", []),
                "comment": payload.get("comment", ""),
                "createdAt": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            }
            data = json.dumps(row_payload, ensure_ascii=False).encode("utf-8")
            request = urllib.request.Request(
                webhook_url,
                data=data,
                headers={"Content-Type": "application/json; charset=utf-8"},
                method="POST",
            )
            try:
                with urllib.request.urlopen(request, timeout=20) as response:
                    response_text = response.read().decode("utf-8", errors="replace")
                self.send_json({"ok": True, "googleResponse": response_text})
            except urllib.error.HTTPError as error:
                error_text = error.read().decode("utf-8", errors="replace")
                self.send_json({"ok": False, "error": f"Google вернул ошибку {error.code}: {error_text}"}, 502)
            except Exception as error:
                self.send_json({"ok": False, "error": str(error)}, 502)
            return

        if path == "/api/company-card":
            payload = self.read_json_payload()
            slug = safe_name((payload.get("company") or {}).get("companyName") or datetime.now().strftime("%Y%m%d%H%M%S"))
            pdf_path = OUTPUTS / f"company_card_{slug}.pdf"
            build_company_card_pdf(payload, pdf_path)
            self.send_json({"pdfUrl": f"/outputs/proposals/{pdf_path.name}"})
            return

        if path != "/api/proposal":
            self.send_error(404)
            return
        payload = self.read_json_payload()
        slug = safe_name(payload.get("number", datetime.now().strftime("%Y%m%d%H%M%S")))
        docx_path = OUTPUTS / f"kp_{slug}.docx"
        pdf_path = OUTPUTS / f"kp_{slug}.pdf"
        oksana_docx_path = OUTPUTS / f"raschet_oksana_{slug}.docx"
        oksana_xlsx_path = OUTPUTS / f"raschet_oksana_{slug}.xlsx"
        build_docx(payload, docx_path)
        build_pdf(payload, pdf_path)
        build_oksana_docx(payload, oksana_docx_path)
        build_oksana_xlsx(payload, oksana_xlsx_path)
        self.send_json({
            "docxUrl": f"/outputs/proposals/{docx_path.name}",
            "pdfUrl": f"/outputs/proposals/{pdf_path.name}",
            "oksanaDocxUrl": f"/outputs/proposals/{oksana_docx_path.name}",
            "oksanaXlsxUrl": f"/outputs/proposals/{oksana_xlsx_path.name}",
        })

    def translate_path(self, path):
        path = unquote(path.split("?", 1)[0].split("#", 1)[0])
        relative = path.lstrip("/") or "index.html"
        return str(ROOT / relative)

    def end_headers(self):
        self.send_header("Cache-Control", "no-store")
        super().end_headers()


if __name__ == "__main__":
    mimetypes.add_type("application/vnd.openxmlformats-officedocument.wordprocessingml.document", ".docx")
    mimetypes.add_type("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", ".xlsx")
    port = int(os.environ.get("PORT", "4173"))
    server = ThreadingHTTPServer(("127.0.0.1", port), Handler)
    print(f"Сайт открыт: http://127.0.0.1:{port}")
    server.serve_forever()
